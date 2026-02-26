"""Conversor de documentos — extrai texto de formatos que o Gemini não lê nativamente.

Formatos que o Gemini lê direto (enviar como bytes):
  - PDF, PNG, JPG, GIF, WEBP, TXT, CSV

Formatos que precisam de conversão para texto:
  - XLSX/XLS → texto tabular (CSV-like)
  - DOCX → texto plano
  - HTML → texto limpo
"""
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Extensões que o Gemini aceita nativamente (enviar bytes direto)
NATIVE_EXTENSIONS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp",
    ".txt", ".csv", ".json", ".xml", ".md",
}

# Extensões que precisam de conversão
CONVERTIBLE_EXTENSIONS = {".xlsx", ".xls", ".docx", ".html", ".htm"}

# Todas as extensões aceitas
SUPPORTED_EXTENSIONS = NATIVE_EXTENSIONS | CONVERTIBLE_EXTENSIONS


def needs_conversion(filename: str) -> bool:
    """Verifica se o arquivo precisa ser convertido para texto."""
    ext = Path(filename).suffix.lower()
    return ext in CONVERTIBLE_EXTENSIONS


def is_supported(filename: str) -> bool:
    """Verifica se o formato é suportado."""
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def convert_to_text(content: bytes, filename: str) -> str:
    """Converte o conteúdo do arquivo para texto legível.

    Args:
        content: bytes do arquivo
        filename: nome do arquivo (para determinar o tipo)

    Returns:
        Texto extraído do documento
    """
    ext = Path(filename).suffix.lower()

    if ext in (".xlsx", ".xls"):
        return _convert_excel(content, filename)
    elif ext == ".docx":
        return _convert_docx(content, filename)
    elif ext in (".html", ".htm"):
        return _convert_html(content)
    else:
        return content.decode("utf-8", errors="replace")


def extract_text_from_pdf(content: bytes, max_pages: int = 50) -> str:
    """Extrai texto de PDF usando pdfplumber com fallback para pypdf."""
    text = _try_pdfplumber(content, max_pages)
    if text and len(text.strip()) > 50:
        return text

    # Fallback
    text_fallback = _try_pypdf(content, max_pages)
    if text_fallback and len(text_fallback.strip()) > len((text or "").strip()):
        return text_fallback

    return text or ""


# ─── Conversores internos ────────────────────────────────────────────

def _convert_excel(content: bytes, filename: str) -> str:
    """Converte Excel (XLSX/XLS) para texto tabular."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Planilha: {sheet_name} ===\n")

            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                # Pula linhas completamente vazias
                if any(cells):
                    rows.append(cells)

            if not rows:
                parts.append("(vazia)\n")
                continue

            # Calcula largura das colunas para formatação tabular
            col_widths = _calc_col_widths(rows, max_width=40)

            for row in rows:
                formatted = " | ".join(
                    str(cell)[:w].ljust(w) for cell, w in zip(row, col_widths)
                )
                parts.append(formatted)

            parts.append("")  # linha em branco entre planilhas

        wb.close()
        result = "\n".join(parts)
        logger.info(f"Excel convertido: {filename} → {len(result)} chars")
        return result

    except Exception as e:
        logger.error(f"Erro ao converter Excel {filename}: {e}")
        return f"[Erro ao converter {filename}: {e}]"


def _convert_docx(content: bytes, filename: str) -> str:
    """Converte DOCX para texto plano preservando estrutura."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Preserva headings com marcação
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style:
                parts.append(f"\n# {text}")
            elif "heading 2" in style:
                parts.append(f"\n## {text}")
            elif "heading 3" in style:
                parts.append(f"\n### {text}")
            else:
                parts.append(text)

        # Tabelas
        for table in doc.tables:
            parts.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
            parts.append("")

        result = "\n".join(parts)
        logger.info(f"DOCX convertido: {filename} → {len(result)} chars")
        return result

    except Exception as e:
        logger.error(f"Erro ao converter DOCX {filename}: {e}")
        return f"[Erro ao converter {filename}: {e}]"


def _convert_html(content: bytes) -> str:
    """Remove tags HTML e retorna texto limpo."""
    import re
    text = content.decode("utf-8", errors="replace")
    # Remove scripts e styles
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Remove tags
    text = re.sub(r"<[^>]+>", " ", text)
    # Limpa espaços
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _try_pdfplumber(content: bytes, max_pages: int) -> str:
    """Extrai texto de PDF com pdfplumber."""
    try:
        import pdfplumber

        pages_text = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages[:max_pages]):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"--- Página {i + 1} ---\n{text}")
        return "\n\n".join(pages_text)

    except Exception as e:
        logger.warning(f"pdfplumber falhou: {e}")
        return ""


def _try_pypdf(content: bytes, max_pages: int) -> str:
    """Extrai texto de PDF com pypdf (fallback)."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages_text = []
        for i, page in enumerate(reader.pages[:max_pages]):
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(f"--- Página {i + 1} ---\n{text}")
        return "\n\n".join(pages_text)

    except Exception as e:
        logger.warning(f"pypdf falhou: {e}")
        return ""


def _calc_col_widths(rows: list[list[str]], max_width: int = 40) -> list[int]:
    """Calcula largura ideal de cada coluna para formatação tabular."""
    if not rows:
        return []
    num_cols = max(len(r) for r in rows)
    widths = [0] * num_cols
    for row in rows:
        for i, cell in enumerate(row):
            widths[i] = max(widths[i], min(len(str(cell)), max_width))
    return [max(w, 3) for w in widths]  # mínimo 3 chars
